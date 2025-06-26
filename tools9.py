# tools9.py
from typing import Union, Dict, List, Tuple
import streamlit as st
import re
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import os
import json
from PyPDF2 import PdfReader
import docx
import oracledb
from langchain_community.vectorstores.oraclevs import OracleVS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from dotenv import load_dotenv
# Load environment variables from .env file
load_dotenv()

def process_document(file):
    """Extract text from various document formats"""
    text = ""
    if file.type == "application/pdf":
        pdf_reader = PdfReader(file)
        text = "".join([page.extract_text() for page in pdf_reader.pages])
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
    else:
        text = file.getvalue().decode()
    return text

def rag_search(query: str) -> str:
    """Search indexed documents for information"""
    if st.session_state.vector_store is None:
        return "No documents have been indexed yet."
    docs = st.session_state.vector_store.similarity_search(query, k=5)
    print("------------in rag_search ")
    print(" the number of pages is " + str(len(docs)))
    return "\n".join([doc.page_content for doc in docs])

def send_email_function(data: Union[str, Dict]) -> str:
    """
    Send email using SMTP. Handles all validation and processing internally.
    This is the actual email sending function separated from the agent tool.
    """
    print(f"Original input: {repr(data)}")
    try:
        if isinstance(data, str):
            print(f"Original input: {repr(data)}")
            cleaned_input = data.strip()
            json_start = cleaned_input.find('{')
            json_end = cleaned_input.rfind('}') + 1
            if json_start != -1 and json_end != -1:
                cleaned_input = cleaned_input[json_start:json_end]
            print(f"Cleaned input: {repr(cleaned_input)}")
            try:
                data = json.loads(cleaned_input)
            except json.JSONDecodeError as e:
                print(f"JSON parsing error: {str(e)}")
                return f"Invalid JSON format: {str(e)}"
        else:
            data = data

        required_fields = {'to', 'subject', 'message'}
        if not all(field in data for field in required_fields):
            missing = required_fields - set(data.keys())
            return f"Missing required fields: {missing}"

        recipients = data['to']
        if isinstance(recipients, str):
            recipients = [r.strip() for r in recipients.split(',')]
        elif isinstance(recipients, list):
            recipients = [r.strip() for r in recipients]
        else:
            return "'to' field must be string or list"

        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        invalid_emails = [email for email in recipients if not re.match(email_pattern, email)]
        if invalid_emails:
            return f"Invalid email addresses: {invalid_emails}"

        msg = MIMEMultipart()
        msg["From"] = os.getenv("SENDER_EMAIL")
        msg["To"] = ", ".join(recipients)
        msg["Subject"] = data['subject'].strip()
        msg.attach(MIMEText(data['message'].strip(), "plain"))

        print(f"SMTP Server: {os.getenv('SMTP_SERVER', 'smtp.gmail.com')}")
        print(f"SMTP Port: {os.getenv('SMTP_PORT', '587')}")
        print(f"Sender Email: {os.getenv('SENDER_EMAIL')}")
        print(f"Sender Password: {'*' * len(os.getenv('SENDER_PASSWORD', ''))} (masked)")
        print(f"Recipients: {msg['To']}")
        print(f"Subject: {msg['Subject']}")
        print(f"Message Body:\n{data['message'].strip()}")

        with smtplib.SMTP(os.getenv("SMTP_SERVER", "smtp.gmail.com"),
                         int(os.getenv("SMTP_PORT", "587"))) as server:
            print("Connecting to SMTP server...")
            server.starttls()
            print("TLS started")
            server.login(os.getenv("SENDER_EMAIL"), os.getenv("SENDER_PASSWORD"))
            print("Logged in successfully")
            server.send_message(msg)
            print("Email sent")

        return f"Email sent successfully to {msg['To']}!"

    except Exception as e:
        print(f"Error details: {str(e)}")
        return f"Error sending email: {str(e)}"

class DatabaseOperations:
    def __init__(self):
        self.connection = None

    def connect(self):
        """Establish connection to Oracle database"""
        try:
            self.connection = oracledb.connect(
                user=os.getenv("ORACLE_USER"),
                password=os.getenv("ORACLE_PASSWORD"),
                dsn=os.getenv("ORACLE_DSN")
            )
            return True
        except Exception as e:
            print(f"Error connecting to database: {e}")
            return False

    def query_recipients(self, name_query: str) -> List[Tuple[str, str, str]]:
        """
        Search for recipients by first name or last name
        Returns a list of tuples containing (first_name, last_name, email)
        """
        try:
            if not self.connection:
                if not self.connect():
                    return []

            cursor = self.connection.cursor()
            search_terms = name_query.strip().split()
            base_query = """
                SELECT first_name, last_name, email
                FROM recipients
                WHERE 1=0
            """
            conditions = []
            params = []

            for term in search_terms:
                conditions.extend([
                    "LOWER(first_name) LIKE LOWER(:term)||'%'",
                    "LOWER(last_name) LIKE LOWER(:term)||'%'"
                ])
                params.extend([term, term])

            query = base_query.replace("1=0", " OR ".join(conditions))

            cursor.execute(query, params)
            results = cursor.fetchall()

            cursor.close()
            return results

        except Exception as e:
            print(f"Error querying recipients: {e}")
            return []

    def close(self):
        """Close the database connection"""
        if self.connection:
            try:
                self.connection.close()
            except Exception as e:
                print(f"Error closing connection: {e}")

def fetch_recipients(query: str) -> str:
    """
    Search for recipients by name and return formatted results
    """
    # Remove any trailing 'O' characters that might be added by the agent
    cleaned_query = query.strip()
    while cleaned_query.endswith('O'):
        cleaned_query = cleaned_query[:-1]

    print(f"Original query: {query}, Cleaned query: {cleaned_query}")

    db_ops = DatabaseOperations()
    recipients = db_ops.query_recipients(cleaned_query)

    if not recipients:
        return f"No recipients found matching the search criteria for '{cleaned_query}'."

    formatted_results = []
    for first_name, last_name, email in recipients:
        formatted_results.append(f"{first_name} {last_name} ({email})")

    # For single recipient, return with suggestion
    if len(recipients) == 1:
        first_name, last_name, email = recipients[0]
        return f"{first_name} {last_name} ({email})\n\nSuggested recipient: {email}"

    # For multiple recipients, make it easy for the agent to choose the first one
    return "\n".join(formatted_results) + "\n\nMultiple recipients found. Using the first email address: " + recipients[0][2]

def extract_email_data_from_response(response_text):
    """
    Extract potential email data from a response.
    Returns a dict with to, subject, and message if found.
    """
    if not response_text:
        return None

    # Try to extract JSON if present (this is the most reliable method)
    json_pattern = r'\{[\s\S]*?\}'
    json_matches = re.findall(json_pattern, response_text)

    for json_str in json_matches:
        try:
            data = json.loads(json_str)
            if all(k in data for k in ["to", "subject", "message"]):
                return data
        except:
            continue

    # Look for patterns that might indicate email data with explicit headers
    to_match = re.search(r"To:\s*([^\n]+)", response_text, re.IGNORECASE)
    subject_match = re.search(r"Subject:\s*([^\n]+)", response_text, re.IGNORECASE)

    # If we found basic email headers, extract the message content
    if to_match and subject_match:
        to = to_match.group(1).strip()
        subject = subject_match.group(1).strip()

        # Extract message content - everything after "Message:" or similar
        message_match = re.search(r"Message:?\s*\n([\s\S]+)$", response_text, re.IGNORECASE)
        if message_match:
            message_content = message_match.group(1).strip()
        else:
            # Try to find content after subject if there's no explicit Message: label
            parts = response_text.split(subject_match.group(0), 1)
            if len(parts) > 1:
                message_content = parts[1].strip()
            else:
                message_content = "No message content found."

        return {
            "to": to,
            "subject": subject,
            "message": message_content
        }

    # Look for email addresses in the response
    # This might be useful if the AI mentioned an email in a less structured way
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    email_match = re.search(email_pattern, response_text)

    if email_match:
        email = email_match.group(0)

        # Look for possible subject nearby (within 200 chars)
        email_pos = response_text.find(email)
        subject_context = response_text[max(0, email_pos-100):min(len(response_text), email_pos+100)]

        subject = "Information"  # Default subject
        for phrase in ["about", "regarding", "concerning", "on the topic of", "subject"]:
            subject_pattern = f"{phrase}\\s+([^.,!?\\n]+)"
            subject_match = re.search(subject_pattern, subject_context, re.IGNORECASE)
            if subject_match:
                subject = subject_match.group(1).strip()
                break

        # Use the rest of the text as message, removing any common prefixes
        message_lines = []
        capture = False
        for line in response_text.split('\n'):
            if email in line:
                capture = True
                continue
            if capture and line.strip():
                message_lines.append(line)

        message = "\n".join(message_lines) if message_lines else response_text

        return {
            "to": email,
            "subject": subject,
            "message": message
        }

    return None

def extract_email_from_text2(text):
    """Extract email address from text that might contain name and email in parentheses"""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(email_pattern, text)
    if match:
        return match.group(0)
    return None

# Tool for use by the agent
def send_email_tool(input_data):
    """
    Process email data from agent and store it in session state
    for the manual email/PDF form
    """
    try:
        # Parse input data
        if isinstance(input_data, str):
            # Try to extract JSON from the input
            try:
                # Look for a JSON object within the string
                json_pattern = r'\{[\s\S]*?\}'
                json_matches = re.findall(json_pattern, input_data)

                if json_matches:
                    for json_str in json_matches:
                        try:
                            data = json.loads(json_str)
                            if all(k in data for k in ["to", "subject", "message"]):
                                break
                        except:
                            continue

                # If no valid JSON found, try to extract data using regex
                if 'data' not in locals():
                    to_match = re.search(r"[Tt]o\s*:\s*([^,\n]+)", input_data)
                    subject_match = re.search(r"[Ss]ubject\s*:\s*([^,\n]+)", input_data)
                    message_match = re.search(r"[Mm]essage\s*:\s*([\s\S]+)$", input_data)

                    if to_match and subject_match and message_match:
                        data = {
                            "to": to_match.group(1).strip(),
                            "subject": subject_match.group(1).strip(),
                            "message": message_match.group(1).strip()
                        }
            except Exception as e:
                print(f"Error parsing input: {str(e)}")
                # If all parsing fails, use a default structure
                data = {
                    "to": "user@example.com",
                    "subject": "Information",
                    "message": input_data
                }
        else:
            # Input is already a dictionary
            data = input_data

        # Extract email address from recipients if it contains name (firstname lastname (email@domain.com))
        #data['message'] = data['message'].replace('\n', '\\n')  # escape for JSON safety

        if "to" in data and "(" in data["to"] and ")" in data["to"]:
            email_match = re.search(r'\((.*?)\)', data["to"])
            if email_match:
                data["to"] = email_match.group(1)

        # Store the data in session state for the form
        st.session_state.email_data = {
            "to": data.get("to", ""),
            "subject": data.get("subject", ""),
            "message": data.get("message", "")
        }

        return "Email content prepared. Please review and choose to either 'Send Email' or 'Save as PDF' from the form."

    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"Error processing email data: {str(e)}"

def extract_email_from_text1(text):
    """Extract email address from text that might contain name and email in parentheses"""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(email_pattern, text)
    if match:
        return match.group(0)
    return None


def send_email_tool_safe(input_data):
    """
    Process email data from agent and store it in session state
    for the manual email/PDF form.
    Safer version with JSON fallback, message escaping, and no crash.
    """
    try:
        data = None  # always define early

        # Handle string input
        if isinstance(input_data, str):
            json_pattern = r'\{[\s\S]*?\}'
            json_matches = re.findall(json_pattern, input_data)

            # Try parsing JSON blocks first
            for json_str in json_matches:
                try:
                    parsed = json.loads(json_str)
                    if isinstance(parsed, dict):
                        parsed['message'] = parsed.get('message', '').replace('\n', '\\n')
                        if all(k in parsed for k in ["to", "subject", "message"]):
                            data = parsed
                            break
                except:
                    continue

            # If JSON failed or was incomplete, fallback to regex
            if data is None:
                to_match = re.search(r"[Tt]o\s*:\s*([^\n]+)", input_data)
                subject_match = re.search(r"[Ss]ubject\s*:\s*([^\n]+)", input_data)
                message_match = re.search(r"[Mm]essage\s*:\s*([\s\S]+)", input_data)

                data = {
                    "to": to_match.group(1).strip() if to_match else "user@example.com",
                    "subject": subject_match.group(1).strip() if subject_match else "Info",
                    "message": message_match.group(1).strip() if message_match else input_data
                }

                data['message'] = data.get('message', '').replace('\n', '\\n')

        # Handle dictionary input
        elif isinstance(input_data, dict):
            data = input_data
            data['message'] = data.get('message', '').replace('\n', '\\n')

        else:
            return "Unsupported input type for send_email_tool."

        # Clean up "to" field if it has (email@example.com) format
        if "to" in data and "(" in data["to"] and ")" in data["to"]:
            email_match = re.search(r'\((.*?)\)', data["to"])
            if email_match:
                data["to"] = email_match.group(1)

        # Save cleaned data to session state
        st.session_state.email_data = {
            "to": data.get("to", ""),
            "subject": data.get("subject", ""),
            "message": data.get("message", "")
        }

        return "Email content prepared. Please review and choose to either 'Send Email' or 'Save as PDF' from the form."

    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"Error processing email data: {str(e)}"

def chunks_to_docs_wrapper(row: dict) -> Document:
    """Converts a row into a Document object for Oracle Vector Store."""
    return Document(page_content=row['text'], metadata={'id': str(row['id']), 'link': row['link']})